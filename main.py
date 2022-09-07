from datetime import datetime
from pathlib import Path
from msvcrt import getch
from os import path

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Cm
from PIL import Image
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor, Pt

number_of_cakes = 4

allergens_dict = {1: 'Glutén', 2: 'Rák', 3: 'Tojás', 4: 'Hal', 5: 'Földimogyoró', 6: 'Szója', 7: 'Tejtermék, laktóz',
                  8: 'Diófélék', 9: 'Zeller', 10: 'Mustár', 11: 'Szezám', 12: 'Kéndioxid szulfit', 13: 'Csillagfürt',
                  14: 'Puhatestűek', 15: 'Mesterséges édesítőszer', 16: 'Édesgyökér'}

months = {1: 'Január', 2: 'Február', 3: 'Március', 4: 'Április', 5: 'Május', 6: 'Június', 7: 'Július', 8: 'Augusztus',
          9: 'Szeptember', 10: 'Október', 11: 'November', 12: 'December'}

url = "https://pupicake.hu/rendeles"
base_url = "https://pupicake.hu/"
user_agent = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:104.0) Gecko/20100101 Firefox/104.0"}


def build_database():
    """returns a dictionary of all the cakes found on pupicake.
    dictionary has cake name, description, link to image, calories and allergens
    dict keys:  img_href
                description
                calories
                allergens

    :return: dictionary
    """

    r = requests.get(url, headers=user_agent)
    soup = BeautifulSoup(r.text, "html.parser")

    confectionery_items = soup.find_all("div", {"class": "confectionery_item"})

    cakes = {}
    for item in confectionery_items:
        cake = item.find_all("a")[1]
        cake_name = cake["data-caption"]
        cake_img_href = base_url + cake["href"]
        cake_info = item.find("div", {"class": "confectionery_item_info"}).text
        info_list = cake_info.strip().split("\n")
        description = info_list[1].replace("\t", "")
        calories_and_allergens = info_list[2].split(", allergen: ")
        calories = calories_and_allergens[0].strip()
        allergens = calories_and_allergens[1].split(",")
        allergens_int = [int(allergen) for allergen in allergens]
        cakes[cake_name] = {"img_href": cake_img_href, "description": description,
                            "calories": calories, "allergens": allergens_int}

    return cakes


def image_to_jpg(image_path):
    path = Path(image_path)
    if path.suffix not in {'.jpg', '.png', '.jfif', '.exif', '.gif', '.tiff', '.bmp'}:
        jpg_image_path = f'{path.parent / path.stem}.jpeg'
        Image.open(image_path).convert('RGB').save(jpg_image_path)
        return jpg_image_path
    return image_path


def build_document(cakes, chosen_cakes):
    assert isinstance(cakes, dict)
    assert isinstance(chosen_cakes, list)

    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    month = datetime.now().month
    heading = document.add_heading(f"Glownexus tortázás 2022 {months[month]}", 0)
    heading.alignment = 1

    table = document.add_table(rows=number_of_cakes, cols=2)

    allergens_in_chosen_cakes = set()
    for index, cake in enumerate(chosen_cakes):
        for allergen in cakes[cake]["allergens"]:
            allergens_in_chosen_cakes.add(allergen)

        row = table.rows[index].cells
        font = row[0].add_paragraph().add_run(cake).font
        font.size = Pt(16)
        font.bold = True
        font.color.rgb = RGBColor(50, 50, 255)
        row[0].add_paragraph(cakes[cake]["description"])
        row[0].add_paragraph(cakes[cake]["calories"])
        row[0].add_paragraph(f"Allergének: {str(cakes[cake]['allergens'])[1:-1]}")
        r = requests.get(cakes[cake]["img_href"], headers=user_agent)
        with open(f"./imgs/{cake}.jpeg", "wb") as f:
            f.write(r.content)
        image = Image.open(f"./imgs/{cake}.jpeg")
        width, height = image.size
        if width / height != 1.5:
            left = 0
            right = width
            new_height = width / 1.5
            difference = height - new_height
            top = difference / 2
            bottom = height - difference / 2
            image = image.crop((left, top, right, bottom))
            image.save(f"./imgs/{cake}.jpeg")

        row[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row[1].aligment = 1
        p = row[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_to_jpg(f"./imgs/{cake}.jpeg"), width=Cm(6), height=Cm(4))

    document.add_paragraph()
    allergens_in_chosen_cakes = sorted(list(allergens_in_chosen_cakes))
    allergens_paragraph = [f"{allergen}: {allergens_dict[allergen]}" for allergen in allergens_in_chosen_cakes]
    allergens_paragraph = str(allergens_paragraph)[1:-1].replace("'", "")
    p = document.add_paragraph(allergens_paragraph)
    p.alignment = 1
    document.save(path.join("output", f"Glownexus tortázás 2022 {months[month]}.docx"))


def cake_input(index, cakes, chosen_cakes):
    assert isinstance(index, int)
    assert isinstance(cakes, dict)
    assert isinstance(chosen_cakes, list)

    cake = input(f"{index}. torta: ")
    if cake in cakes:
        chosen_cakes.append(cake)
    else:
        print("Cake not found")
        return cake_input(index, cakes=cakes, chosen_cakes=chosen_cakes)

if __name__ == "__main__":
    cakes = build_database()

    chosen_cakes = []
    for index in range(1, 5):
        cake_input(index=index, cakes=cakes, chosen_cakes=chosen_cakes)

    print("Starting to build document.")
    build_document(cakes=cakes, chosen_cakes=chosen_cakes)
    print("All done, press any key to exit.")
    _ = getch()
